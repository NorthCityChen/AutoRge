'''
@Author: Mr.Sen
@LastEditTime: 2020-05-19 23:23:40
@Website: 
@Mr.Sen All rights reserved
'''
from docx import Document
from docx.shared import Inches,Pt
import code_pic
import random
import json
import os,shutil
import init

def main():
    init.init()
    cd=code_pic.main()
    # cd={'5812':'woca'}
    with open(os.getcwd()+'/settings/codesave.json','w') as f:
        json.dump(cd,f)
        # cd=json.load(f)
    with open(os.getcwd()+'/settings/zjie.json','r',encoding='utf-8') as f2:
        # print(type(f2))
        content=f2.read()
        if content.startswith(u'\ufeff'):
            content=content.encode('utf8')[3:].decode('utf8')
        py=json.loads(content)
        # py=eval(content)
    for i in cd:
        shutil.copyfile('mb.docx',os.getcwd()+"\\bg\\"+str(i)+'.docx')
    for i in cd:
        doc=Document(os.getcwd()+"\\bg\\"+str(i)+'.docx')
        table=doc.tables[0]
        table.cell(1,4).text=cd.get(i)
        # table.cell(1,6).text=py[str(random.randint(1,3))]
        table.cell(1,6).text=random.choice(list(py.values()))
        run=table.cell(1,0).paragraphs[0].add_run()
        run.add_picture(os.getcwd()+"\\pic\\"+str(i)+'.png',width=Inches(4.0),height=Inches(2.5))
        doc.paragraphs[10].add_run("%s" % i).font.size=Pt(20)
        doc.paragraphs[11].add_run("%s" % i).font.size=Pt(20)
        doc.save(os.getcwd()+"\\bg\\"+str(i)+'.docx')
if __name__=='__main__':
    main()